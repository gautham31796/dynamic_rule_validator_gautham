import pandas as pd
import json
import fitz  # PyMuPDF
import docx
import re
import operator as op

def extract_text_from_pdf(pdf_path):
    with fitz.open(pdf_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_word(doc_path):
    doc = docx.Document(doc_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def normalize_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    return ' '.join(text.split())

def find_paragraph_with_text(doc_path, target_text):
    doc = docx.Document(doc_path)
    target_text_clean = normalize_text(target_text)
    for para in doc.paragraphs:
        full_text = ''.join([run.text for run in para.runs])
        if normalize_text(full_text).find(target_text_clean) != -1:
            return para
    return None

def clean_font_name(font_name):
    return re.sub(r'[^a-z]', '', font_name.lower())

def validate_style(paragraph, style_requirements):
    style_req = style_requirements.lower()
    required_font = None
    required_size = None
    required_bold = False

    if "style:" in style_req:
        try:
            required_font = style_req.split("style:")[1].split()[0].strip().lower()
        except:
            pass
    if "size:" in style_req:
        try:
            required_size = float(style_req.split("size:")[1].split()[0])
        except:
            pass
    if "bold" in style_req:
        required_bold = True

    font_match = size_match = bold_match = False

    for run in paragraph.runs:
        run_font = run.font.name or ""
        run_size = run.font.size.pt if run.font.size else None
        run_bold = run.bold

        if required_font and required_font in clean_font_name(run_font):
            font_match = True
        if required_size is not None and run_size is not None and abs(run_size - required_size) < 0.5:
            size_match = True
        if required_bold and run_bold:
            bold_match = True

    if required_font and not font_match:
        return False, "Font mismatch"
    if required_size is not None and not size_match:
        return False, "Size mismatch"
    if required_bold and not bold_match:
        return False, "Bold mismatch"

    return True, "Style matched"

def validate_pdf_style(pdf_path, expected_text, style_requirements):
    doc = fitz.open(pdf_path)
    expected_norm = normalize_text(expected_text)
    style_req = style_requirements.lower()
    required_font = None
    required_bold = False
    required_size = None

    if "style:" in style_req:
        try:
            required_font = style_req.split("style:")[1].split()[0].strip().lower()
        except:
            pass
    if "bold" in style_req:
        required_bold = True
    if "size:" in style_req:
        try:
            required_size = float(style_req.split("size:")[1].split()[0])
        except:
            pass

    for page in doc:
        all_spans = []
        all_text = ""
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    all_spans.append((text, span))
                    all_text += f"{text} "

        if expected_norm in normalize_text(all_text):
            for text, span in all_spans:
                norm_span_text = normalize_text(text)
                if norm_span_text in expected_norm or expected_norm.startswith(norm_span_text):
                    font_name = span.get("font", "").lower()
                    font_size = span.get("size", 0)
                    is_bold = "bold" in font_name or (span.get("flags", 0) & 2 != 0)

                    font_match = size_match = bold_match = False

                    if required_font and required_font in clean_font_name(font_name):
                        font_match = True
                    if required_size and abs(font_size - required_size) <= 0.5:
                        size_match = True
                    if required_bold and is_bold:
                        bold_match = True

                    if (not required_font or font_match) and (required_size is None or size_match) and (not required_bold or bold_match):
                        return True, "Style matched"
                    else:
                        return False, f"Style mismatch: font_match={font_match}, size_match={size_match}, bold_match={bold_match}"
            return False, "Text matched, but no span matched for style validation"

    return False, "Expected text not found in PDF for style validation"

def evaluate_rule(rule_row, document_text, input_data, document_path):
    rule_id = rule_row.get('Output Identifier', 'N/A')
    input_val = rule_row.get('Input Value', '')
    expected = rule_row['Output Language']
    style_req = rule_row.get('Style', '').strip()

    input_data_lower = {k.lower(): v for k, v in input_data.items()}
    all_conditions = [cond.strip() for cond in re.split(r'\n|;', input_val) if cond.strip()]

    OPERATORS = {
        '=': op.eq,
        '<>': op.ne,
        '>': op.gt,
        '<': op.lt,
        '>=': op.ge,
        '<=': op.le
    }

    def parse_condition(condition):
        for symbol in ['<>', '>=', '<=', '>', '<', '=']:
            if symbol in condition:
                key, val = condition.split(symbol, 1)
                return key.strip().lower(), symbol, val.strip().strip('"').strip("'")
        return None, None, None

    for cond in all_conditions:
        key, symbol, expected_val = parse_condition(cond)
        if not symbol or key not in input_data_lower:
            continue

        actual_val = input_data_lower[key]

        if isinstance(actual_val, dict):
            actual_val = list(actual_val.keys())[0] if actual_val else ""

        if isinstance(actual_val, list):
            norm_list = [normalize_text(str(v)) for v in actual_val]
            expected_vals = [normalize_text(v.strip()) for v in expected_val.split(',')]
            if symbol == '=' and not all(val in norm_list for val in expected_vals):
                return 'SKIPPED', f"{expected_vals} not all found in list for {key}"
            elif symbol == '<>' and any(val in norm_list for val in expected_vals):
                return 'SKIPPED', f"Some values {expected_vals} unexpectedly found in list for {key}"
            continue

        try:
            actual_float = float(actual_val)
            expected_float = float(expected_val)
            if not OPERATORS[symbol](actual_float, expected_float):
                return 'SKIPPED', f"Condition Mismatch: {key} {symbol} {expected_val} failed (numeric)"
        except ValueError:
            actual_norm = normalize_text(str(actual_val))
            expected_norm = normalize_text(expected_val)
            if not OPERATORS[symbol](actual_norm, expected_norm):
                return 'SKIPPED', f"Condition Mismatch: {key} {symbol} {expected_val} failed (text)"

    placeholders = re.findall(r"<(.*?)>", expected)
    for key in placeholders:
        if key not in input_data:
            return 'FAIL', f"Missing input data for placeholder <{key}>"
        val = input_data.get(key, "")
        if isinstance(val, dict):
            val = list(val.values())[0] if val else ""
        expected = expected.replace(f"<{key}>", str(val))

    expected_clean = normalize_text(expected)
    document_text_clean = normalize_text(document_text)

    if len(expected_clean) < 3:
        return 'FAIL', "Expected text too short after placeholder replacement"

    if expected_clean in document_text_clean:
        if style_req:
            if document_path.lower().endswith('.docx'):
                para = find_paragraph_with_text(document_path, expected)
                if para:
                    style_ok, style_reason = validate_style(para, style_req)
                    if not style_ok:
                        return 'FAIL', f"Style validation failed — {style_reason}"
                else:
                    return 'FAIL', "Text matched but paragraph not found for style validation"
            elif document_path.lower().endswith('.pdf'):
                style_ok, style_reason = validate_pdf_style(document_path, expected, style_req)
                if not style_ok:
                    return 'FAIL', f"PDF Style validation failed — {style_reason}"
        return 'PASS', "All conditions met and text matched"
    else:
        return 'FAIL', "Expected output not found in document"