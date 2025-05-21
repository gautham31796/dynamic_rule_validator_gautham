import pandas as pd
import json
import fitz  # PyMuPDF
import docx
import re

def extract_text_from_pdf(pdf_path):
    with fitz.open(pdf_path) as doc:
        return "".join([page.get_text() for page in doc])

def extract_text_from_word(doc_path):
    doc = docx.Document(doc_path)
    return "\n".join([para.text for para in doc.paragraphs])

def normalize_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    return ' '.join(text.split())

def find_paragraph_with_text(doc_path, target_text):
    doc = docx.Document(doc_path)
    target_text_clean = normalize_text(target_text)
    for para in doc.paragraphs:
        full_text = ''.join([run.text for run in para.runs])
        if normalize_text(full_text).find(target_text_clean) != -1:
            return para, doc
    return None, doc

def validate_style(paragraph, style_requirements, doc):
    style_req = style_requirements.lower()
    required_font = None
    required_size = None
    required_bold = False

    if "style:" in style_req:
        required_font = style_req.split("style:")[1].split()[0].strip().lower()
    if "size:" in style_req:
        required_size = float(style_req.split("size:")[1].split()[0])
    if "bold" in style_req:
        required_bold = True

    for run in paragraph.runs:
        font_name = run.font.name if run.font and run.font.name else None
        font_size = run.font.size.pt if run.font and run.font.size else None
        is_bold = run.bold

        if not font_size:
            try:
                style_id = paragraph.style.style_id
                style = doc.styles[style_id]
                if style and style.font.size:
                    font_size = style.font.size.pt
            except:
                pass

        font_match = size_match = bold_match = True

        if required_font:
            clean_font = re.sub(r'[^a-z]', '', (font_name or "").lower())
            font_match = required_font in clean_font

        if required_size is not None:
            size_match = font_size is not None and abs(font_size - required_size) < 0.5

        if required_bold:
            bold_match = bool(is_bold)

        if font_match and size_match and bold_match:
            return True, "Style matched"

    return False, "Style mismatch"

def validate_pdf_style(pdf_path, expected_text, style_requirements):
    doc = fitz.open(pdf_path)
    expected_norm = normalize_text(expected_text)
    style_req = style_requirements.lower()

    required_font = None
    required_bold = False
    required_size = None

    if "style:" in style_req:
        required_font = style_req.split("style:")[1].split()[0].strip().lower()
    if "bold" in style_req:
        required_bold = True
    if "size:" in style_req:
        required_size = float(style_req.split("size:")[1].split()[0])

    for page in doc:
        spans = [span for block in page.get_text("dict")["blocks"]
                      for line in block.get("lines", [])
                      for span in line.get("spans", [])]
        all_text = " ".join(span["text"] for span in spans)

        if expected_norm in normalize_text(all_text):
            for span in spans:
                text = span["text"]
                if normalize_text(text) in expected_norm:
                    font_name = span.get("font", "").lower()
                    font_size = span.get("size", 0)
                    is_bold = "bold" in font_name or (span.get("flags", 0) & 2 != 0)

                    font_match = size_match = bold_match = True

                    if required_font and required_font not in re.sub(r'[^a-z]', '', font_name):
                        font_match = False
                    if required_size and abs(font_size - required_size) > 0.5:
                        size_match = False
                    if required_bold and not is_bold:
                        bold_match = False

                    if font_match and size_match and (not required_bold or bold_match):
                        return True, "Style matched"
                    else:
                        return False, "PDF style mismatch"

            return False, "Text matched, no styled span matched"

    return False, "Expected text not found in PDF"

def evaluate_rule(rule_row, document_text, input_data, document_path):
    input_val = rule_row.get('Input Value', '')
    expected = rule_row['Output Language']
    style_req = rule_row.get('Style', '').strip()
    input_data_lower = {k.lower(): v for k, v in input_data.items()}
    all_conditions = [cond.strip() for cond in re.split(r'\n|;', input_val) if cond.strip()]

    for cond in all_conditions:
        if '=' not in cond:
            continue
        key, val = cond.split('=', 1)
        key = key.strip().lower()
        expected_val = str(val).strip()
        actual_val = input_data_lower.get(key, "")

        if isinstance(actual_val, dict):
            actual_val = list(actual_val.keys())[0] if actual_val else ""

        expected_values = re.findall(r'"([^"]+)"', expected_val) or [v.strip() for v in expected_val.split(',')]
        expected_values = [normalize_text(v) for v in expected_values]

        if isinstance(actual_val, list):
            actual_norm_list = [normalize_text(str(v)) for v in actual_val]
            if any(val not in actual_norm_list for val in expected_values):
                return 'SKIPPED', f"List Mismatch for {key}"
        else:
            actual_norm = normalize_text(str(actual_val))
            if actual_norm != expected_values[0]:
                return 'SKIPPED', f"Condition Mismatch for {key}"

    for key in re.findall(r"<(.*?)>", expected):
        val = input_data.get(key, "")
        if isinstance(val, dict):
            val = list(val.values())[0] if val else ""
        expected = expected.replace(f"<{key}>", str(val))

    if normalize_text(expected) in normalize_text(document_text):
        if style_req:
            if document_path.lower().endswith('.docx'):
                para, doc = find_paragraph_with_text(document_path, expected)
                if para:
                    style_ok, style_reason = validate_style(para, style_req, doc)
                    if not style_ok:
                        return 'FAIL', style_reason
            elif document_path.lower().endswith('.pdf'):
                style_ok, style_reason = validate_pdf_style(document_path, expected, style_req)
                if not style_ok:
                    return 'FAIL', style_reason
        return 'PASS', "Validation passed"
    else:
        return 'FAIL', "Expected output not found"

def load_rules(excel_path):
    df = pd.read_excel(excel_path, engine='openpyxl')
    df.columns = df.columns.str.strip()
    return df

def main(rule_path, doc_path, json_path, output_path):
    rules_df = load_rules(rule_path)

    if doc_path.lower().endswith('.pdf'):
        document_text = extract_text_from_pdf(doc_path)
    elif doc_path.lower().endswith('.docx'):
        document_text = extract_text_from_word(doc_path)
    else:
        raise ValueError("Unsupported document type")

    with open(json_path, 'r') as f:
        raw_data = json.load(f)
        input_data = raw_data.get("testData", raw_data)

    results = []
    for _, row in rules_df.iterrows():
        result, reason = evaluate_rule(row, document_text, input_data, doc_path)
        results.append({
            "Output Identifier": row.get("Output Identifier"),
            "Status": result,
            "Reason": reason
        })

    pd.DataFrame(results).to_excel(output_path, index=False)
