import fitz
import docx

def parse_document(file_path):
    if file_path.lower().endswith(".pdf"):
        return parse_pdf(file_path)
    elif file_path.lower().endswith(".docx"):
        return parse_word(file_path)
    else:
        raise ValueError("Unsupported file type")

def parse_pdf(path):
    doc = fitz.open(path)
    paragraphs = []
    for page_num, page in enumerate(doc, 1):
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block['type'] == 0:
                for line in block["lines"]:
                    text = " ".join(span["text"] for span in line["spans"])
                    paragraphs.append({"text": text, "page": page_num})
    return {"paragraphs": paragraphs, "tables": []}

def parse_word(path):
    doc = docx.Document(path)
    paragraphs = [{"text": p.text} for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        table_data = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(table_data)
    return {"paragraphs": paragraphs, "tables": tables}
