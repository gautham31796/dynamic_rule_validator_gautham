import pandas as pd
import json
import fitz  # PyMuPDF
import docx
import re
import operator

def extract_text_from_pdf(pdf_path):
    with fitz.open(pdf_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_word(doc_path):
    doc = docx.Document(doc_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def normalize_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    return ' '.join(text.split())

def find_paragraph_with_text(doc_path, target_text):
    doc = docx.Document(doc_path)
    target_text_clean = normalize_text(target_text)
    for para in doc.paragraphs:
        full_text = ''.join([run.text for run in para.runs])
        if normalize_text(full_text).find(target_text_clean) != -1:
            return para
    return None

def validate_style(paragraph, style_requirements):
    style_req = style_requirements.lower()
    font_match = size_match = bold_match = False

    for run in paragraph.runs:
        font = run.font
        if not font:
            continue

        font_name = font.name.lower() if font.name else ""
        font_size = font.size.pt if font.size else None
        is_bold = run.bold

        if "times new roman" in style_req and "times new roman" in font_name:
            font_match = True
        if "bold" in style_req and is_bold:
            bold_match = True
        if "size:" in style_req and font_size:
            try:
                expected_size = float(style_req.split("size:")[1].split()[0])
                if abs(font_size - expected_size) < 0.2:
                    size_match = True
            except:
                pass

    all_requirements = []
    if "times new roman" in style_req:
        all_requirements.append(font_match)
    if "bold" in style_req:
        all_requirements.append(bold_match)
    if "size:" in style_req:
        all_requirements.append(size_match)

    if all(all_requirements):
        return True, "Style matched"
    else:
        return False, f"Style mismatch: font_match={font_match}, bold_match={bold_match}, size_match={size_match}"

def validate_pdf_style(pdf_path, expected_text, style_requirements):
    doc = fitz.open(pdf_path)
    style_req = style_requirements.lower()
    expected_norm = normalize_text(expected_text)

    for page in doc:
        all_spans = []
        all_text = ""

        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    all_spans.append((text, span))
                    all_text += f"{text} "

        all_text_norm = normalize_text(all_text)

        if expected_norm in all_text_norm:
            for text, span in all_spans:
                if normalize_text(text) in expected_norm:
                    font_name = span.get("font", "")
                    font_size = span.get("size", 0)

                    if "times new roman" in style_req and "times" not in font_name.lower():
                        return False, f"Font mismatch: got '{font_name}'"
                    if "bold" in style_req and "bold" not in font_name.lower():
                        return False, f"Bold style missing in font '{font_name}'"
                    if "size:" in style_req:
                        try:
                            expected_size = float(style_req.split("size:")[1].split()[0])
                            if abs(font_size - expected_size) > 0.5:
                                return False, f"Font size mismatch: got {font_size}, expected {expected_size}"
                        except:
                            pass
                    return True, "Style matched"
            return True, "Text matched, but could not confirm styling span — assuming pass"

    return False, "Text not found in PDF for style validation"

def evaluate_rule(rule_row, document_text, input_data, document_path):
    input_val = rule_row.get('Input Value', '')
    expected = rule_row['Output Language']
    style_req = rule_row.get('Style', '').strip()

    input_data_lower = {k.lower(): v for k, v in input_data.items()}
    all_conditions = [cond.strip() for cond in input_val.split('\n') if cond.strip()]

    OPERATORS = {
        '=': operator.eq,
        '<>': operator.ne,
        '>': operator.gt,
        '<': operator.lt
    }

    def parse_condition(condition):
        for op in ['<>', '>=', '<=', '>', '<', '=']:
            if op in condition:
                key, val = map(str.strip, condition.split(op, 1))
                return key.lower(), op, val.strip().strip('"').strip("'")
        return None, None, None

    for cond in all_conditions:
        key, op, expected_val = parse_condition(cond)
        if not op or key not in input_data_lower:
            continue

        actual_val = input_data_lower[key]

        if isinstance(actual_val, list):
            norm_list = [normalize_text(str(v)) for v in actual_val]
            expected_vals = [normalize_text(v.strip()) for v in expected_val.split(',')]

            if op == '=' and not all(val in norm_list for val in expected_vals):
                return 'SKIPPED', f"{expected_vals} not all found in list for {key}"
            elif op == '<>' and any(val in norm_list for val in expected_vals):
                return 'SKIPPED', f"Some values {expected_vals} unexpectedly found in list for {key}"
            continue

        actual_val = str(actual_val).strip().strip('"').strip("'")
        try:
            actual_val_float = float(actual_val)
            expected_val_float = float(expected_val)
            compare_result = OPERATORS[op](actual_val_float, expected_val_float)
        except ValueError:
            compare_result = OPERATORS[op](normalize_text(actual_val), normalize_text(expected_val))

        if not compare_result:
            return 'SKIPPED', f"Condition Mismatch: {key} {op} {expected_val} failed"

    placeholders = re.findall(r"<(.*?)>", expected)
    for key in placeholders:
        val = input_data.get(key, "")
        if isinstance(val, dict):
            val = list(val.values())[0] if val else ""
        expected = expected.replace(f"<{key}>", str(val))

    expected_clean = normalize_text(expected)
    document_text_clean = normalize_text(document_text)

    if expected_clean in document_text_clean:
        if style_req:
            if document_path.lower().endswith('.docx'):
                para = find_paragraph_with_text(document_path, expected)
                if para:
                    style_ok, style_reason = validate_style(para, style_req)
                    if not style_ok:
                        return 'FAIL', f"Style validation failed — {style_reason}"
                else:
                    return 'FAIL', "Text matched but paragraph not found for style validation"
            elif document_path.lower().endswith('.pdf'):
                style_ok, style_reason = validate_pdf_style(document_path, expected, style_req)
                if not style_ok:
                    return 'FAIL', f"PDF Style validation failed — {style_reason}"
        return 'PASS', "All conditions met and text matched"
    else:
        return 'FAIL', "Expected output not found in document"