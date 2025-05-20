import pandas as pd
import json
import fitz  # PyMuPDF
import docx
import re

def extract_text_from_pdf(pdf_path):
    with fitz.open(pdf_path) as doc:
        text = ""
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_word(doc_path):
    doc = docx.Document(doc_path)
    return '\n'.join([para.text for para in doc.paragraphs])

def normalize_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    return ' '.join(text.split())

def find_paragraph_with_text(doc_path, target_text):
    doc = docx.Document(doc_path)
    target_text_clean = normalize_text(target_text)
    for para in doc.paragraphs:
        full_text = ''.join([run.text for run in para.runs])
        if normalize_text(full_text).find(target_text_clean) != -1:
            return para
    return None

def clean_font_name(font_name):
    return re.sub(r'[^a-z]', '', font_name.lower())

def validate_style(paragraph, style_requirements):
    style_req = style_requirements.lower()
    required_font = None
    required_size = None
    required_bold = False

    if "style:" in style_req:
        try:
            required_font = style_req.split("style:")[1].split()[0].strip().lower()
        except:
            pass
    if "size:" in style_req:
        try:
            required_size = float(style_req.split("size:")[1].split()[0])
        except:
            pass
    if "bold" in style_req:
        required_bold = True

    for run in paragraph.runs:
        font_name = None
        font_size = None

        if run.font and run.font.name:
            font_name = run.font.name
        elif run.style and run.style.font and run.style.font.name:
            font_name = run.style.font.name
        elif paragraph.style and paragraph.style.font and paragraph.style.font.name:
            font_name = paragraph.style.font.name

        if run.font and run.font.size:
            font_size = run.font.size.pt
        elif run.style and run.style.font and run.style.font.size:
            font_size = run.style.font.size.pt
        elif paragraph.style and paragraph.style.font and paragraph.style.font.size:
            font_size = paragraph.style.font.size.pt

        is_bold = run.bold

        print(f"[DEBUG] Run text: '{run.text.strip()}' | font: {font_name} | size: {font_size} | bold: {is_bold}")

        font_match = size_match = bold_match = True

        if required_font:
            if font_name:
                clean_font = re.sub(r'[^a-z]', '', font_name.lower())
                font_match = required_font in clean_font
            else:
                font_match = False

        if required_size is not None:
            size_match = font_size is not None and abs(font_size - required_size) < 0.5

        if required_bold:
            bold_match = bool(is_bold)

        if font_match and size_match and bold_match:
            return True, "Style matched"

    return False, f"Style mismatch: font_match={font_match}, size_match={size_match}, bold_match={bold_match}"

def validate_pdf_style(pdf_path, expected_text, style_requirements):
    import re

    doc = fitz.open(pdf_path)
    expected_norm = normalize_text(expected_text)
    style_req = style_requirements.lower()

    required_font = None
    required_bold = False
    required_size = None

    if "style:" in style_req:
        try:
            required_font = style_req.split("style:")[1].split()[0].strip().lower()
        except:
            pass
    if "bold" in style_req:
        required_bold = True
    if "size:" in style_req:
        try:
            required_size = float(style_req.split("size:")[1].split()[0])
        except:
            pass

    for page in doc:
        all_spans = []
        all_text = ""

        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    all_spans.append((text, span))
                    all_text += f"{text} "

        # Text must be found somewhere in the page
        if expected_norm in normalize_text(all_text):
            for text, span in all_spans:
                norm_span_text = normalize_text(text)
                if norm_span_text in expected_norm or expected_norm.startswith(norm_span_text):
                    font_name = span.get("font", "").lower()
                    font_size = span.get("size", 0)
                    is_bold = "bold" in font_name or (span.get("flags", 0) & 2 != 0)

                    print(f"[DEBUG] PDF Span matched: '{text}' | font: '{font_name}', size: {font_size}, bold: {is_bold}")

                    font_match = size_match = bold_match = True

                    if required_font and required_font not in re.sub(r'[^a-z]', '', font_name):
                        font_match = False
                    if required_size and abs(font_size - required_size) > 0.5:
                        size_match = False
                    if required_bold and not is_bold:
                        bold_match = False

                    if font_match and size_match and (not required_bold or bold_match):
                        return True, "Style matched"
                    else:
                        return False, f"Style mismatch: font_match={font_match}, size_match={size_match}, bold_match={bold_match}"

            # Text found but no matching styled span
            return False, "Text matched, but no span matched for style validation"

    return False, "Expected text not found in PDF for style validation"

def evaluate_rule(rule_row, document_text, input_data, document_path):
    rule_id = rule_row.get('Output Identifier', 'N/A')
    input_val = rule_row.get('Input Value', '')
    expected = rule_row['Output Language']
    style_req = rule_row.get('Style', '').strip()

    print(f"\n[DEBUG] Evaluating Rule {rule_id}")
    print(f"[DEBUG] Raw input_val = {repr(input_val)}")

    input_data_lower = {k.lower(): v for k, v in input_data.items()}

    all_conditions = [cond.strip() for cond in re.split(r'\n|;', input_val) if cond.strip()]
    print(f"[DEBUG] Parsed conditions ({len(all_conditions)}):")
    for cond in all_conditions:
        print(f"  -> {cond}")

    for cond in all_conditions:
        if '=' not in cond:
            continue
        key, val = cond.split('=', 1)
        key = key.strip().lower()
        expected_val = str(val).strip()

        actual_val = input_data_lower.get(key, "")

        if isinstance(actual_val, dict):
            actual_val = list(actual_val.keys())[0] if actual_val else ""

        expected_values = re.findall(r'"([^"]+)"', expected_val) or [v.strip() for v in expected_val.split(',')]
        expected_values = [normalize_text(v) for v in expected_values]

        if isinstance(actual_val, list):
            actual_norm_list = [normalize_text(str(v)) for v in actual_val]
            print(f"[DEBUG] Compare list — key: {key}, expected: {expected_values}, actual: {actual_norm_list}")
            missing = [val for val in expected_values if val not in actual_norm_list]
            if missing:
                return 'SKIPPED', f"List Mismatch for {key}: missing values {missing}"
        else:
            actual_norm = normalize_text(str(actual_val))
            print(f"[DEBUG] Compare single — key: {key}, expected: {expected_values[0]}, actual: {actual_norm}")
            if len(expected_values) > 1:
                return 'SKIPPED', f"Expected multiple values for {key} but field is not a list"
            if actual_norm != expected_values[0]:
                return 'SKIPPED', f"Condition Mismatch for {key}: expected '{expected_values[0]}', got '{actual_norm}'"

    # Replace placeholders in expected text
    placeholders = re.findall(r"<(.*?)>", expected)
    for key in placeholders:
        val = input_data.get(key, "")
        if isinstance(val, dict):
            val = list(val.values())[0] if val else ""
        expected = expected.replace(f"<{key}>", str(val))

    expected_clean = normalize_text(expected)
    document_text_clean = normalize_text(document_text)

    print(f"[DEBUG] FINAL TEXT MATCH CHECK — expected: {expected_clean[:80]}...")

    if expected_clean in document_text_clean:
        if style_req:
            if document_path.lower().endswith('.docx'):
                para = find_paragraph_with_text(document_path, expected)
                if para:
                    style_ok, style_reason = validate_style(para, style_req)
                    if not style_ok:
                        return 'FAIL', f"Style validation failed — {style_reason}"
                else:
                    return 'FAIL', "Text matched but paragraph not found for style validation"
            elif document_path.lower().endswith('.pdf'):
                style_ok, style_reason = validate_pdf_style(document_path, expected, style_req)
                if not style_ok:
                    return 'FAIL', f"PDF Style validation failed — {style_reason}"
        return 'PASS', "All conditions met and text matched"
    else:
        return 'FAIL', "Expected output not found in document"

def load_rules(excel_path):
    df = pd.read_excel(excel_path, engine='openpyxl')
    df.columns = df.columns.str.strip()
    return df

def main():
    excel_path = "Rules.xlsx"
    document_path = "1_of_1_GAI1356789_AccidentInsurance_GroupCertificate_EC1.docx"  # or .pdf
    json_path = "testdata.json"

    rules_df = load_rules(excel_path)

    if document_path.lower().endswith('.pdf'):
        document_text = extract_text_from_pdf(document_path)
    elif document_path.lower().endswith('.docx'):
        document_text = extract_text_from_word(document_path)
    else:
        raise ValueError("Unsupported document type. Use PDF or Word (.docx)")

    with open(json_path, 'r') as f:
        raw_data = json.load(f)
        input_data = raw_data.get("testData", raw_data)

    output_data = []

    for _, row in rules_df.iterrows():
        print(f"\n[INFO] Running rule: {row.get('Output Identifier')}")
        result, reason = evaluate_rule(row, document_text, input_data, document_path)
        output_data.append({
            "Output Identifier": row.get('Output Identifier'),
            "Status": result,
            "Reason": reason
        })
        print(f"[RESULT] Rule {row.get('Output Identifier')}: {result} — {reason}")

    result_df = pd.DataFrame(output_data)
    result_df.to_excel("rule_results.xlsx", index=False)

if __name__ == "__main__":
    main()
